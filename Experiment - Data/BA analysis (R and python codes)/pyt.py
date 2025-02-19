from statsmodels.formula.api import ols
import pandas as pd
from statsmodels.stats.outliers_influence import variance_inflation_factor
from docx import Document

from docx import Document

print("Debugging VIF Data...")


# Replace 'path_to_your_dataset.csv' with the actual path to your CSV file
data_cleaned = pd.read_csv('/Users/monikakaczorowska/Desktop/DSBA/semester 1/Business Analytics/project/BA analysis/dataCleaned.csv')


# Step 1: Define the regression models
# Define the formulas for each model
model_1_formula = 'Trust_Change ~ Treatment_Randomizer'
model_2_formula = model_1_formula + ' + pre_trust_segment'
model_3_formula = model_2_formula + ' + Tech_Adoption_Openess'
model_4_formula = model_3_formula + ' + Treatment_Randomizer:pre_trust_segment'
model_5_formula = model_4_formula + ' + AI_Usage_Frequency + AI_Knowledge_Rating + Digital_Privacy_Concern + pre_confidence_segment'

# Fit the models
model_1 = ols(model_1_formula, data=data_cleaned).fit()
model_2 = ols(model_2_formula, data=data_cleaned).fit()
model_3 = ols(model_3_formula, data=data_cleaned).fit()
model_4 = ols(model_4_formula, data=data_cleaned).fit()
model_5 = ols(model_5_formula, data=data_cleaned).fit()

# Collect models into a list
models = [model_1, model_2, model_3, model_4, model_5]
model_names = ['Model 1', 'Model 2', 'Model 3', 'Model 4', 'Model 5']

# Step 2: Format results for all models
def format_coef_pval(coef, pval):
    return f"{coef:.4f}{'*' if pval < 0.05 else ''}"  # Add asterisk for significance

formatted_results = pd.DataFrame()
for model_name, model in zip(model_names, models):
    model_data = {
        var: format_coef_pval(coef, model.pvalues.get(var, float('nan')))
        for var, coef in model.params.items()
    }
    model_df = pd.DataFrame.from_dict(model_data, orient='index', columns=[model_name])
    formatted_results = pd.concat([formatted_results, model_df], axis=1)

# Reorder variables to match the desired order for display
variable_order = [
    "Intercept", "Treatment_Randomizer[T.Negative]", "Treatment_Randomizer[T.Positive]",
    "pre_trust_segment[T.Medium]", "pre_trust_segment[T.High]",
    "Tech_Adoption_Openess",
    "Treatment_Randomizer[T.Negative]:pre_trust_segment[T.Medium]",
    "Treatment_Randomizer[T.Positive]:pre_trust_segment[T.Medium]",
    "Treatment_Randomizer[T.Negative]:pre_trust_segment[T.High]",
    "Treatment_Randomizer[T.Positive]:pre_trust_segment[T.High]",
    "AI_Usage_Frequency", "AI_Knowledge_Rating", "Digital_Privacy_Concern",
    "pre_confidence_segment[T.Medium]", "pre_confidence_segment[T.High]"
]
formatted_results = formatted_results.reindex(variable_order).reset_index()
formatted_results.columns = ['Variable'] + model_names



# Step 3: Perform Variance Inflation Factor (VIF) analysis for Model 5
vif_data = data_cleaned[['Treatment_Randomizer', 'pre_trust_segment', 'Tech_Adoption_Openess',
                         'AI_Usage_Frequency', 'AI_Knowledge_Rating', 'Digital_Privacy_Concern',
                         'pre_confidence_segment']]


vif_data = pd.get_dummies(vif_data, drop_first=True)  # Convert categorical to dummy variables


print(vif_data.head())  # Display the first few rows
print(vif_data.dtypes)  # Check the data types of all columns
print(vif_data.isnull().sum())  # Check for missing values

vif_values = pd.DataFrame({
    "Variable": vif_data.columns,
    "VIF": [variance_inflation_factor(vif_data.values, i) for i in range(vif_data.shape[1])]
})

# Step 4: Add R-squared and Adjusted R-squared values to the formatted results
rsquared_values = [model.rsquared for model in models]
adjusted_rsquared_values = [model.rsquared_adj for model in models]

rsquared_data = pd.DataFrame({
    "Variable": ["R-squared", "Adjusted R-squared"],
    **{model_name: [f"{rsq:.4f}", f"{adj_rsq:.4f}"] for model_name, rsq, adj_rsq in zip(model_names, rsquared_values, adjusted_rsquared_values)}
})

formatted_results_with_rsquared = pd.concat([formatted_results, rsquared_data], ignore_index=True)

# Step 5: Save regression results and VIF analysis to a Word document
doc = Document()
doc.add_heading('Regression Results and Analysis', level=1)

# Add regression results table to the document
doc.add_heading('Regression Results', level=2)
table = doc.add_table(rows=1, cols=formatted_results_with_rsquared.shape[1])
table.style = 'Table Grid'

# Add headers
header_cells = table.rows[0].cells
for i, column_name in enumerate(formatted_results_with_rsquared.columns):
    header_cells[i].text = column_name

# Add rows
for index, row in formatted_results_with_rsquared.iterrows():
    cells = table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = str(value)

# Add VIF analysis to the document
doc.add_heading('VIF Analysis for Model 5', level=2)
vif_table = doc.add_table(rows=1, cols=vif_values.shape[1])
vif_table.style = 'Table Grid'

# Add headers for VIF table
vif_header_cells = vif_table.rows[0].cells
for i, column_name in enumerate(vif_values.columns):
    vif_header_cells[i].text = column_name

# Add rows for VIF table
for index, row in vif_values.iterrows():
    vif_cells = vif_table.add_row().cells
    for i, value in enumerate(row):
        vif_cells[i].text = str(value)

# Save the document
word_file_path = '/mnt/data/Regression_Results_and_VIF.docx'
doc.save(word_file_path)

word_file_path
